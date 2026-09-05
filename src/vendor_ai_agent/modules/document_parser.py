"""Visual-semantic document parser using pdfplumber."""
from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Iterable, List, Optional, Tuple

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image
except ImportError:
    pytesseract = None
    convert_from_path = None
    Image = None

from ..contracts import DocumentParserContract
from ..models import DocExtracted, TenderSection
from .document_processing import DocumentClassifier, FieldExtractor, SectionExtractor


class DocumentParser(DocumentParserContract):
    """Converts tender files into structured sections using visual analysis."""

    HEADER_KEYWORDS = [
        "SCOPE", "REQUIREMENT", "DELIVERABLE", "MANDATORY", "EVALUATION",
        "AMENDMENT", "ADDENDUM", "SECTION", "PART", "ARTICLE", "CLAUSE",
        "SCHEDULE", "APPENDIX", "ANNEX", "TABLE", "SPECIFICATION",
        "GENERAL", "SPECIFIC", "TECHNICAL", "PERFORMANCE", "QUALITY"
    ]

    def __init__(self) -> None:
        self.logger = logging.getLogger(__name__)
        self.classifier = DocumentClassifier()
        self.section_extractor = SectionExtractor()
        self.field_extractor = FieldExtractor()

    def parse(self, files: Iterable[Path]) -> List[TenderSection]:
        """Parse files into structured sections using visual-semantic analysis."""
        collected_paths = self._collect_paths(files)
        sections: List[TenderSection] = []
        
        for path in collected_paths:
            parser = self._select_parser(path)
            if parser is None:
                self.logger.debug("Unsupported file type for %s", path)
                continue
            
            try:
                parsed = parser(path)
                doc_type = self.classifier.classify(path).doc_type
                for section in parsed:
                    section.metadata["doc_type"] = doc_type.value
                sections.extend(parsed)
            except Exception as exc:
                self.logger.error("Failed to parse %s: %s", path, exc)
                sections.append(
                    TenderSection(
                        title=path.stem,
                        content="",
                        source_path=path,
                        section_type="error",
                        metadata={"error": str(exc)},
                    )
                )
        return sections

    def _collect_paths(self, targets: Iterable[Path]) -> List[Path]:
        paths: List[Path] = []
        for target in targets:
            if target.is_dir():
                for nested in target.rglob("*"):
                    if nested.is_file():
                        paths.append(nested)
            elif target.is_file():
                paths.append(target)
        return paths

    def _select_parser(self, path: Path):
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._parse_pdf
        if suffix in {".xlsx", ".xls"}:
            return self._parse_excel
        if suffix in {".docx", ".doc"}:
            return self._parse_docx
        if suffix in {".txt", ".md", ".csv"}:
            return self._parse_text
        return None

    def _parse_pdf(self, path: Path) -> List[TenderSection]:
        """Parse PDF using pdfplumber with visual anchor detection."""
        if pdfplumber is None:
            raise RuntimeError("pdfplumber is required but not installed")
        
        sections: List[TenderSection] = []
        current_section_title = path.stem
        current_buffer: List[str] = []
        
        with pdfplumber.open(path) as pdf:
            base_font_size = self._detect_base_font_size(pdf)
            
            for page_num, page in enumerate(pdf.pages, 1):
                tables = page.extract_tables()
                
                if tables:
                    if current_buffer:
                        sections.append(
                            TenderSection(
                                title=current_section_title,
                                content="\n".join(current_buffer).strip(),
                                source_path=path,
                                section_type="text",
                                metadata={"file_type": "pdf", "page": page_num},
                            )
                        )
                        current_buffer = []
                    
                    for idx, table in enumerate(tables, 1):
                        table_md = self._table_to_markdown(table)
                        if table_md:
                            sections.append(
                                TenderSection(
                                    title=f"{current_section_title} (Table {idx})",
                                    content=table_md,
                                    source_path=path,
                                    section_type="table",
                                    metadata={"file_type": "pdf", "page": page_num, "table_index": idx},
                                )
                            )
                
                text = page.extract_text()
                is_ocr = False
                
                if not text or len(text.strip()) < 50:
                    ocr_text = self._try_ocr_on_page(path, page_num)
                    if ocr_text:
                        text = ocr_text
                        is_ocr = True
                        self.logger.info(f"OCR applied to page {page_num} of {path.name}")
                    else:
                        continue
                
                lines = text.split('\n')
                
                if is_ocr:
                    words_data = []
                else:
                    words_data = page.extract_words(extra_attrs=["fontname", "size"])
                
                for line in lines:
                    line_stripped = line.strip()
                    
                    if not line_stripped or len(line_stripped) < 3:
                        current_buffer.append(line)
                        continue
                    
                    if self._is_header_or_footer(line_stripped):
                        continue
                    
                    if self._is_likely_header(line_stripped, words_data, base_font_size):
                        if current_buffer:
                            sections.append(
                                TenderSection(
                                    title=current_section_title,
                                    content="\n".join(current_buffer).strip(),
                                    source_path=path,
                                    section_type="text",
                                    metadata={"file_type": "pdf"},
                                )
                            )
                            current_buffer = []
                        
                        current_section_title = f"{path.stem} – {line_stripped}"
                    else:
                        current_buffer.append(line)
        
        if current_buffer:
            sections.append(
                TenderSection(
                    title=current_section_title,
                    content="\n".join(current_buffer).strip(),
                    source_path=path,
                    section_type="text",
                    metadata={"file_type": "pdf"},
                )
            )
        
        if not sections:
            full_text = "\n".join(current_buffer).strip()
            sections.append(
                TenderSection(
                    title=path.stem,
                    content=full_text,
                    source_path=path,
                    section_type="text",
                    metadata={"file_type": "pdf"},
                )
            )
        
        return sections

    def _detect_base_font_size(self, pdf) -> float:
        """Detect the most common font size (base text size)."""
        font_sizes = {}
        
        for page in pdf.pages[:min(3, len(pdf.pages))]:
            words = page.extract_words(extra_attrs=["size"])
            for word in words:
                size = round(word.get('size', 12), 1)
                font_sizes[size] = font_sizes.get(size, 0) + 1
        
        if not font_sizes:
            return 12.0
        
        return max(font_sizes.items(), key=lambda x: x[1])[0]

    def _is_likely_header(self, line: str, words_data: List[dict], base_size: float) -> bool:
        """Detect if line is a section header using multiple heuristics."""
        if len(line) > 150:
            return False
        
        if any(kw in line.upper() for kw in self.HEADER_KEYWORDS):
            if len(line) < 100:
                return True
        
        if re.match(r'^(SECTION|PART|ARTICLE|CLAUSE)?\s*\d+(\.\d+)*', line, re.IGNORECASE):
            return True
        
        if line.isupper() and 3 < len(line.split()) <= 12:
            return True
        
        line_words = [w for w in words_data if w['text'] in line]
        if line_words:
            avg_size = sum(w.get('size', base_size) for w in line_words) / len(line_words)
            if avg_size > base_size + 1:
                return True
        
        return False

    def _is_header_or_footer(self, line: str) -> bool:
        """Detect common header/footer patterns to exclude."""
        patterns = [
            r'^Page\s+\d+\s+of\s+\d+',
            r'^\d+\s*/\s*\d+$',
            r'^RFB\s+tender\s+#\d+',
            r'^PR-\d+-',
            r'^\d{2}-\d{2}-\d{2}$',
        ]
        
        for pattern in patterns:
            if re.match(pattern, line, re.IGNORECASE):
                return True
        
        return False

    def _table_to_markdown(self, table: List[List]) -> str:
        """Convert table to markdown format."""
        if not table or not table[0]:
            return ""
        
        cleaned = []
        for row in table:
            cleaned_row = []
            for cell in row:
                if cell is None:
                    cleaned_row.append('')
                else:
                    cell_str = str(cell).replace('\n', ' ').replace('|', '\\|').strip()
                    cleaned_row.append(cell_str)
            cleaned.append(cleaned_row)
        
        if not cleaned or not cleaned[0]:
            return ""
        
        md = "| " + " | ".join(cleaned[0]) + " |\n"
        md += "| " + " | ".join(['---'] * len(cleaned[0])) + " |\n"
        
        for row in cleaned[1:]:
            md += "| " + " | ".join(row) + " |\n"
        
        return md

    def _try_ocr_on_page(self, pdf_path: Path, page_num: int) -> Optional[str]:
        """Apply OCR to a single PDF page if text extraction failed."""
        if pytesseract is None or convert_from_path is None:
            return None
        
        try:
            images = convert_from_path(
                pdf_path,
                first_page=page_num,
                last_page=page_num,
                dpi=300
            )
            
            if not images:
                return None
            
            text = pytesseract.image_to_string(images[0], lang='eng')
            
            if text and len(text.strip()) > 50:
                return text
            
            return None
            
        except Exception as exc:
            self.logger.warning(f"OCR failed for page {page_num} of {pdf_path.name}: {exc}")
            return None

    def _parse_excel(self, path: Path) -> List[TenderSection]:
        """Parse Excel file as markdown tables."""
        sections: List[TenderSection] = []
        if openpyxl is None:
            raise RuntimeError("openpyxl is required for Excel parsing but not installed")
        
        workbook = openpyxl.load_workbook(path, data_only=True)
        for sheet in workbook.worksheets:
            table_data = []
            for row in sheet.iter_rows(values_only=True):
                cells = ['' if value is None else str(value) for value in row]
                table_data.append(cells)
            
            if table_data:
                md_content = self._table_to_markdown(table_data)
            else:
                md_content = ""
            
            sections.append(
                TenderSection(
                    title=f"{path.stem} – {sheet.title}",
                    content=md_content,
                    source_path=path,
                    section_type="table",
                    metadata={"file_type": "excel", "sheet": sheet.title},
                )
            )
        
        return sections

    def _parse_docx(self, path: Path) -> List[TenderSection]:
        """Parse Word document (.docx) with structure preservation."""
        if Document is None:
            raise RuntimeError("python-docx is required for Word parsing but not installed")
        
        sections: List[TenderSection] = []
        doc = Document(path)
        
        current_section_title = path.stem
        current_buffer: List[str] = []
        
        for element in doc.element.body:
            if element.tag.endswith('p'):
                para = next((p for p in doc.paragraphs if p._element == element), None)
                if para is None:
                    continue
                
                text = para.text.strip()
                if not text:
                    current_buffer.append('')
                    continue
                
                if self._is_word_header(para):
                    if current_buffer:
                        sections.append(
                            TenderSection(
                                title=current_section_title,
                                content="\n".join(current_buffer).strip(),
                                source_path=path,
                                section_type="text",
                                metadata={"file_type": "docx"},
                            )
                        )
                        current_buffer = []
                    
                    current_section_title = f"{path.stem} – {text}"
                else:
                    current_buffer.append(text)
            
            elif element.tag.endswith('tbl'):
                table = next((t for t in doc.tables if t._element == element), None)
                if table is None:
                    continue
                
                if current_buffer:
                    sections.append(
                        TenderSection(
                            title=current_section_title,
                            content="\n".join(current_buffer).strip(),
                            source_path=path,
                            section_type="text",
                            metadata={"file_type": "docx"},
                        )
                    )
                    current_buffer = []
                
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    md_content = self._table_to_markdown(table_data)
                    sections.append(
                        TenderSection(
                            title=f"{current_section_title} (Table)",
                            content=md_content,
                            source_path=path,
                            section_type="table",
                            metadata={"file_type": "docx"},
                        )
                    )
        
        if current_buffer:
            sections.append(
                TenderSection(
                    title=current_section_title,
                    content="\n".join(current_buffer).strip(),
                    source_path=path,
                    section_type="text",
                    metadata={"file_type": "docx"},
                )
            )
        
        if not sections:
            full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            sections.append(
                TenderSection(
                    title=path.stem,
                    content=full_text,
                    source_path=path,
                    section_type="text",
                    metadata={"file_type": "docx"},
                )
            )
        
        return sections

    def _is_word_header(self, para) -> bool:
        """Detect if Word paragraph is a section header."""
        text = para.text.strip()
        
        if len(text) > 150:
            return False
        
        if para.style.name.startswith('Heading'):
            return True
        
        if any(kw in text.upper() for kw in self.HEADER_KEYWORDS):
            if len(text) < 100:
                return True
        
        if re.match(r'^(SECTION|PART|ARTICLE|CLAUSE)?\s*\d+(\.\d+)*', text, re.IGNORECASE):
            return True
        
        if text.isupper() and 3 < len(text.split()) <= 12:
            return True
        
        if para.runs and para.runs[0].bold and len(text.split()) <= 15:
            return True
        
        return False

    def _parse_text(self, path: Path) -> List[TenderSection]:
        """Parse plain text files."""
        content = path.read_text(encoding="utf-8", errors="ignore")
        if path.suffix.lower() == ".md":
            sections = []
            title, lines = path.stem, []
            fence = None
            for line in content.splitlines():
                marker = re.match(r"^\s{0,3}(`{3,}|~{3,})", line)
                if marker:
                    token = marker.group(1)
                    if fence is None:
                        fence = token
                    elif token[0] == fence[0] and len(token) >= len(fence):
                        fence = None
                heading = None if fence or marker else re.match(r"^ {0,3}#{1,6}\s+(.+?)\s*#*\s*$", line)
                if heading:
                    if any(line.strip() for line in lines) or title != path.stem:
                        sections.append(TenderSection(title, "\n".join(lines).strip(), path,
                                                      metadata={"file_type": "markdown"}))
                    title, lines = heading.group(1), []
                else:
                    lines.append(line)
            if any(line.strip() for line in lines) or title != path.stem:
                sections.append(TenderSection(title, "\n".join(lines).strip(), path,
                                              metadata={"file_type": "markdown"}))
            return sections
        return [
            TenderSection(
                title=path.stem,
                content=content.strip(),
                source_path=path,
                section_type="text",
                metadata={"file_type": "text"},
            )
        ]
