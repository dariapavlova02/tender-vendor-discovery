"""Document classifier with content-based detection fallback."""
from __future__ import annotations

import re
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
from typing import Optional

import pdfplumber


class DocumentType(str, Enum):
    CORE_RFP = "core_rfp"
    CORE_SCOPE = "core_scope"
    TECH_SPEC = "tech_spec"
    TECH_AMENDMENT = "tech_amendment"
    ADDENDUM = "addendum"
    APPENDIX = "appendix"
    PRESENTATION = "presentation"
    LEGAL = "legal"
    OTHER = "other"


@dataclass
class ClassifiedDocument:
    path: Path
    doc_type: DocumentType
    doc_number: Optional[int] = None
    priority_score: float = 0.0
    title_hint: Optional[str] = None
    size_bytes: int = 0


class DocumentClassifier:
    """Enhanced classifier with document number extraction and priority scoring."""

    CORE_KEYWORDS = ["solicitation", "rfp", "rfq", "rfb", "sow", "scope", "bid document", "request for"]
    TECH_KEYWORDS = ["spec", "technical", "boq", "schedule", "attachment"]
    ADDENDUM_KEYWORDS = ["addendum", "addenda", "amendment"]
    APPENDIX_KEYWORDS = ["appendix", "schedule", "table"]
    LEGAL_KEYWORDS = ["terms", "agreement", "contract", "legal"]
    PRESENTATION_KEYWORDS = ["presentation", "briefing", "info session"]

    def classify(self, path: Path) -> ClassifiedDocument:
        name_lower = path.name.lower()
        size = path.stat().st_size
        
        doc_number = self._extract_number(name_lower)
        doc_type = self._classify_type(path, name_lower)
        priority_score = self._calculate_priority(doc_type, doc_number, size)
        
        return ClassifiedDocument(
            path=path,
            doc_type=doc_type,
            doc_number=doc_number,
            priority_score=priority_score,
            title_hint=path.stem,
            size_bytes=size
        )
    
    def _extract_number(self, name: str) -> Optional[int]:
        """Extract document number from filename."""
        patterns = [
            r'addendum\s*#(\d+)',
            r'amendment\s*no\.?\s*(\d+)',
            r'addenda\s*#(\d+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, name, re.I)
            if match:
                return int(match.group(1))
        
        return None
    
    def _classify_type(self, path: Path, name_lower: str) -> DocumentType:
        """Classify document type with fine granularity."""
        content_type = self._peek_inside(path)
        if content_type:
            return content_type
        
        if 'amendment' in name_lower:
            if any(kw in name_lower for kw in ['pr-', 'spec', 'requirement', 'technical']):
                return DocumentType.TECH_AMENDMENT
            return DocumentType.ADDENDUM
        
        if any(kw in name_lower for kw in ['addendum', 'addenda']):
            return DocumentType.ADDENDUM
        
        if any(kw in name_lower for kw in self.CORE_KEYWORDS):
            if not any(kw in name_lower for kw in ['addendum', 'amendment']):
                return DocumentType.CORE_RFP
        
        if any(kw in name_lower for kw in self.APPENDIX_KEYWORDS):
            return DocumentType.APPENDIX
        
        if any(kw in name_lower for kw in self.PRESENTATION_KEYWORDS):
            return DocumentType.PRESENTATION
        
        if any(kw in name_lower for kw in self.TECH_KEYWORDS):
            return DocumentType.TECH_SPEC
        
        if any(kw in name_lower for kw in self.LEGAL_KEYWORDS):
            return DocumentType.LEGAL
        
        return DocumentType.OTHER
    
    def _calculate_priority(self, doc_type: DocumentType, doc_number: Optional[int], size: int) -> float:
        """Calculate priority score (higher = process first)."""
        tier_scores = {
            DocumentType.CORE_RFP: 1000,
            DocumentType.CORE_SCOPE: 1000,
            DocumentType.TECH_AMENDMENT: 900,
            DocumentType.TECH_SPEC: 850,
            DocumentType.APPENDIX: 700,
            DocumentType.ADDENDUM: 600,
            DocumentType.PRESENTATION: 400,
            DocumentType.LEGAL: 300,
            DocumentType.OTHER: 100,
        }
        
        score = tier_scores.get(doc_type, 100)
        
        if doc_number:
            score += doc_number * 10
        
        size_bonus = min(size / 1_000_000 * 50, 200)
        score += size_bonus
        
        return score

    def _peek_inside(self, path: Path) -> Optional[DocumentType]:
        """Analyze content to detect document type."""
        suffix = path.suffix.lower()
        
        if suffix == '.pdf':
            return self._peek_pdf(path)
        elif suffix in {'.docx', '.doc'}:
            return self._peek_docx(path)
        
        return None
    
    def _peek_pdf(self, path: Path) -> Optional[DocumentType]:
        """Analyze first 2 pages of PDF."""
        try:
            with pdfplumber.open(path) as pdf:
                pages_to_read = min(2, len(pdf.pages))
                text_sample = ""
                
                for i in range(pages_to_read):
                    page_text = pdf.pages[i].extract_text()
                    if page_text:
                        text_sample += page_text[:1000]
                
                return self._classify_by_content(text_sample)
                
        except Exception:
            return None
    
    def _peek_docx(self, path: Path) -> Optional[DocumentType]:
        """Analyze first paragraphs of Word document."""
        try:
            from docx import Document
            
            doc = Document(path)
            text_sample = ""
            
            for para in doc.paragraphs[:20]:
                text_sample += para.text + "\n"
                if len(text_sample) > 2000:
                    break
            
            return self._classify_by_content(text_sample)
            
        except Exception:
            return None
    
    def _classify_by_content(self, text: str) -> Optional[DocumentType]:
        """Classify document based on text content."""
        if not text:
            return None
        
        text_lower = text.lower()
        
        if 'amendment no' in text_lower and any(kw in text_lower for kw in ['requirement', 'specification', 'shall meet', 'general requirments']):
            return DocumentType.TECH_AMENDMENT
        
        if 'addendum no' in text_lower and 'notice to all' in text_lower:
            return DocumentType.ADDENDUM
        
        if any(kw in text_lower for kw in ["addendum no", "amendment no", "addenda"]):
            return DocumentType.ADDENDUM
        
        if any(kw in text_lower for kw in ["request for proposal", "request for quotation", "rfp", "rfq", "rfb", "solicitation"]):
            return DocumentType.CORE_RFP
        
        if 'appendix' in text_lower and ('table' in text_lower or 'schedule' in text_lower):
            return DocumentType.APPENDIX
        
        if any(kw in text_lower for kw in ["technical specification", "attachment", "schedule a", "schedule b"]):
            return DocumentType.TECH_SPEC
        
        if any(kw in text_lower for kw in ["terms and conditions", "general conditions", "contract agreement"]):
            return DocumentType.LEGAL
        
        return None
